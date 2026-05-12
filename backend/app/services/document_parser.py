import fitz  # PyMuPDF
from docx import Document
from pathlib import Path


class DocumentParser:

    def parse_file(self, file_path: str) -> str:
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return self.parse_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return self.parse_docx(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    def parse_pdf(self, file_path: str) -> str:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text.strip()

    def parse_docx(self, file_path: str) -> str:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)


document_parser = DocumentParser()
